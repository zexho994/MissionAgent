from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "generated" / "harness-core-summary.docx"


COLORS = {
    "ink": "172033",
    "muted": "5B667A",
    "blue": "2E5EAA",
    "teal": "1B8A7A",
    "paper": "F7F4EE",
    "line": "D9E2EF",
}


def set_paragraph_shading(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    run = p.add_run(text)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    run.font.color.rgb = RGBColor.from_string(COLORS["ink"])
    return p


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.18
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r2 = p.add_run(text[len(bold_prefix):])
        runs = [r1, r2]
    else:
        runs = [p.add_run(text)]
    for run in runs:
        run.font.name = "Arial"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor.from_string(COLORS["ink"])
    return p


def add_callout(doc, title, body):
    p = doc.add_paragraph()
    set_paragraph_shading(p, COLORS["paper"])
    p.paragraph_format.left_indent = Cm(0.25)
    p.paragraph_format.right_indent = Cm(0.25)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(12)
    r1 = p.add_run(f"{title}：")
    r1.bold = True
    r2 = p.add_run(body)
    for r in (r1, r2):
        r.font.name = "Arial"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
        r.font.size = Pt(10.5)
        r.font.color.rgb = RGBColor.from_string(COLORS["ink"])


def add_bullet(doc, label, body):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.15
    r1 = p.add_run(f"{label}：")
    r1.bold = True
    r2 = p.add_run(body)
    for run in (r1, r2):
        run.font.name = "Arial"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor.from_string(COLORS["ink"])


def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.1)
    section.right_margin = Cm(2.1)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    styles["Normal"].font.size = Pt(10.5)
    for name, size, color in [
        ("Heading 1", 17, COLORS["blue"]),
        ("Heading 2", 13, COLORS["teal"]),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(6)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("Harness：长任务 Agent 的运行框架")
    run.bold = True
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor.from_string(COLORS["ink"])

    meta = doc.add_paragraph()
    meta_run = meta.add_run("核心摘要 · 基于 Anthropic 2026-03-24 文章理解")
    meta_run.font.name = "Arial"
    meta_run._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    meta_run.font.size = Pt(9)
    meta_run.font.color.rgb = RGBColor.from_string(COLORS["muted"])

    add_callout(
        doc,
        "一句话",
        "Harness 不是某个 agent，而是管理 agent 如何拆任务、交接上下文、执行、验收、复盘的系统框架。",
    )

    add_heading(doc, "为什么重要", 1)
    add_body(doc, "长任务不能只靠一个更强的模型。任务做久后，agent 容易目标漂移、上下文混乱、提前收尾，或者对自己的结果过度乐观。")
    add_body(doc, "Harness 的价值，是把复杂任务变成可管理的流程：有人规划，有人执行，有人验收，有结构化交接。")

    add_heading(doc, "核心结构", 1)
    add_bullet(doc, "Planner / Owner", "把模糊目标拆成清楚的任务和完成标准。")
    add_bullet(doc, "Generator / Worker", "按任务执行，产出可检查的 artifact。")
    add_bullet(doc, "Evaluator / Supervisor", "独立验收，不相信执行者自评。")
    add_bullet(doc, "Handoff Artifact", "上下文重置时，用结构化文件交接状态和下一步。")

    add_heading(doc, "概念边界", 1)
    add_bullet(doc, "Agent", "做事的角色，例如负责人、研究员、内容创作者、审核员。")
    add_bullet(doc, "Harness", "让这些角色长期可靠协作的运行框架。")
    add_bullet(doc, "OpenClaw", "外部动作执行器，负责浏览器、小红书页面和采集工具。")

    add_heading(doc, "对 DigitalAgent 的结论", 1)
    add_body(doc, "我们要做的不是简单的多 agent 聊天系统，而是 Mission Harness。")
    add_body(doc, "DigitalAgent Core 负责目标、团队、任务、记忆、审核和复盘；OpenClaw 只负责浏览器和外部工具执行。")
    add_body(doc, "OwnerAgent 可以决定需要哪些业务角色，HRAgent 创建 subagents，但必须受 RoleSpec、预算、权限和验收标准约束。")

    add_heading(doc, "第一版原则", 1)
    add_body(doc, "先跑通一个小闭环：创建 Mission → 组队 → 分配任务 → OpenClaw 执行 → 回传 artifact → 审核 → 生成下一轮任务。")
    add_body(doc, "不要一开始追求 24 小时全自治。先证明系统能持续产出、被审核、能复盘，再扩大自动化范围。")

    source = doc.add_paragraph()
    source.paragraph_format.space_before = Pt(10)
    r = source.add_run("来源：https://www.anthropic.com/engineering/harness-design-long-running-apps")
    r.font.name = "Arial"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor.from_string(COLORS["muted"])

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
